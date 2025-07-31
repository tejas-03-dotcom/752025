import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, simpledialog
import language_tool_python
from textblob import TextBlob
from docx import Document
import PyPDF2
import os


def read_file(filepath):
    if filepath.endswith(".txt"):
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    elif filepath.endswith(".docx"):
        doc = Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs])
    elif filepath.endswith(".pdf"):
        text = ""
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                ptext = page.extract_text()
                if ptext:
                    text += ptext
        return text
    else:
        messagebox.showerror("Unsupported File", "Only .txt, .docx, and .pdf files supported.")
        return ""


class GrammarSpellChecker:
    def __init__(self, root):
        self.root = root
        root.title("Grammar & Spelling Checker")
        root.geometry("1100x700")
        self.dark_mode = False
        self.current_content = ""
        self.matches = []
        self.tool = language_tool_python.LanguageTool('en-US')
        self.last_export_path = None

        self.build_ui()

    def build_ui(self):
        self.header = tk.Label(self.root, text="Upload Document to Check Grammar and Spelling",
                               font=("Helvetica", 16, "bold"))
        self.header.pack(pady=10)

        btn_frame = tk.Frame(self.root)
        btn_frame.pack(pady=5)

        self.upload_btn = tk.Button(btn_frame, text="📁 Upload Document", bg="#4CAF50", fg="white",
                                    font=("Helvetica", 12), command=self.upload_and_check)
        self.upload_btn.grid(row=0, column=0, padx=5)

        self.auto_correct_btn = tk.Button(btn_frame, text="Auto-Correct All", bg="#2196F3", fg="white",
                                          font=("Helvetica", 12), command=self.auto_correct_all)
        self.auto_correct_btn.grid(row=0, column=1, padx=5)

        self.export_btn = tk.Button(btn_frame, text="Export Corrected", bg="#FF9800", fg="white",
                                    font=("Helvetica", 12), command=self.export_corrected)
        self.export_btn.grid(row=0, column=2, padx=5)

        self.dark_mode_btn = tk.Button(btn_frame, text="Toggle Dark Mode", bg="#555", fg="white",
                                       font=("Helvetica", 12), command=self.toggle_dark_mode)
        self.dark_mode_btn.grid(row=0, column=3, padx=5)

        self.tooltip = tk.Label(self.root, text="", bg="yellow", fg="black", font=("Helvetica", 10), wraplength=400)

        main_frame = tk.Frame(self.root)
        main_frame.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

        self.text_display = scrolledtext.ScrolledText(main_frame, width=90, height=30, wrap=tk.WORD,
                                                      font=("Courier", 12))
        self.text_display.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.text_display.tag_config("error", background="#FFDDDD", foreground="red")

        self.suggestions_list = tk.Listbox(main_frame, width=40, height=30, font=("Helvetica", 10))
        self.suggestions_list.pack(side=tk.RIGHT, padx=10, fill=tk.Y)

        self.text_display.tag_bind("error", "<Button-1>", self.on_click_error)
        self.text_display.tag_bind("error", "<Enter>", self.on_hover_error)
        self.text_display.tag_bind("error", "<Leave>", lambda e: self.hide_tooltip())

        self.tag_match_map = {}

    def upload_and_check(self):
        filepath = filedialog.askopenfilename(filetypes=[
            ("Supported files", "*.txt *.docx *.pdf"),
            ("All files", "*.*")
        ])
        if not filepath:
            return

        self.current_content = read_file(filepath)
        if not self.current_content:
            return

        self.text_display.delete("1.0", tk.END)
        self.suggestions_list.delete(0, tk.END)
        self.text_display.insert(tk.END, self.current_content)
        self.clear_all_tags()

        self.matches = self.tool.check(self.current_content)

        self.tag_match_map.clear()

        for i, match in enumerate(self.matches):
            start = match.offset
            end = start + match.errorLength
            start_index = f"1.0 + {start} chars"
            end_index = f"1.0 + {end} chars"

            tag_name = f"error_{i}"
            self.text_display.tag_add(tag_name, start_index, end_index)
            self.text_display.tag_config(tag_name, background="#FFDDDD", foreground="red")
            self.text_display.tag_bind(tag_name, "<Button-1>", self.on_click_error)
            self.text_display.tag_bind(tag_name, "<Enter>", self.on_hover_error)
            self.text_display.tag_bind(tag_name, "<Leave>", lambda e: self.hide_tooltip())

            self.tag_match_map[tag_name] = match

            self.suggestions_list.insert(tk.END, f"Issue: {match.message}")
            if match.replacements:
                # Show up to 5 suggestions now
                self.suggestions_list.insert(tk.END, f"Suggestion: {', '.join(match.replacements[:5])}")
            self.suggestions_list.insert(tk.END, "—" * 50)

        blob = TextBlob(self.current_content)
        words = blob.words
        for word in words:
            corrected = str(TextBlob(word).correct())
            if corrected.lower() != word.lower():
                self.suggestions_list.insert(tk.END, f"Spelling: {word} → {corrected}")
                self.suggestions_list.insert(tk.END, "—" * 50)

    def clear_all_tags(self):
        for tag in self.text_display.tag_names():
            self.text_display.tag_remove(tag, "1.0", tk.END)

    def on_hover_error(self, event):
        idx = self.text_display.index(f"@{event.x},{event.y}")
        tags = self.text_display.tag_names(idx)
        for tag in tags:
            if tag.startswith("error_") and tag in self.tag_match_map:
                msg = self.tag_match_map[tag].message
                self.show_tooltip(msg)
                break

    def show_tooltip(self, message):
        self.tooltip.config(text=message)
        self.tooltip.place(relx=0.5, rely=0.05, anchor="n")

    def hide_tooltip(self):
        self.tooltip.place_forget()

    def on_click_error(self, event):
        idx = self.text_display.index(f"@{event.x},{event.y}")
        tags = self.text_display.tag_names(idx)
        for tag in tags:
            if tag.startswith("error_") and tag in self.tag_match_map:
                match = self.tag_match_map[tag]
                self.show_correction_popup(match, tag)
                break

    def show_correction_popup(self, match, tag):
        popup = tk.Toplevel(self.root)
        popup.title("Correction Options")
        popup.geometry("400x250")
        popup.transient(self.root)
        popup.grab_set()

        tk.Label(popup, text=f"Issue: {match.message}", wraplength=380, font=("Helvetica", 12, "bold")).pack(pady=10)

        suggestions = match.replacements
        if not suggestions:
            suggestions = ["No suggestions available"]

        var = tk.StringVar(value=suggestions[0])

        for s in suggestions[:5]:
            tk.Radiobutton(popup, text=s, variable=var, value=s, font=("Helvetica", 11)).pack(anchor='w')

        def apply_correction():
            correction = var.get()
            if correction == "No suggestions available":
                popup.destroy()
                return
            start = match.offset
            end = start + match.errorLength
            start_index = f"1.0 + {start} chars"
            end_index = f"1.0 + {end} chars"
            self.text_display.delete(start_index, end_index)
            self.text_display.insert(start_index, correction)

            self.current_content = self.text_display.get("1.0", tk.END)
            popup.destroy()
            self.upload_and_check()

        tk.Button(popup, text="Apply Correction", bg="#4CAF50", fg="white",
                  command=apply_correction).pack(pady=15)

    def auto_correct_all(self):
        content = self.text_display.get("1.0", tk.END)
        new_content = content
        for match in sorted(self.matches, key=lambda m: m.offset, reverse=True):
            if match.replacements:
                start = match.offset
                end = start + match.errorLength
                correction = match.replacements[0]
                new_content = new_content[:start] + correction + new_content[end:]

        self.text_display.delete("1.0", tk.END)
        self.text_display.insert(tk.END, new_content)
        self.current_content = new_content
        self.upload_and_check()

    def export_corrected(self):
        corrected_text = self.text_display.get("1.0", tk.END).strip()
        if not corrected_text:
            messagebox.showwarning("Nothing to Export", "There is no text to export.")
            return

        if self.last_export_path and os.path.exists(os.path.dirname(self.last_export_path)):
            filepath = self.last_export_path
        else:
            filepath = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
            if not filepath:
                return
            self.last_export_path = filepath

        try:
            if filepath.endswith(".txt"):
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(corrected_text)
            elif filepath.endswith(".docx"):
                doc = Document()
                for line in corrected_text.splitlines():
                    doc.add_paragraph(line)
                doc.save(filepath)
            else:
                messagebox.showerror("Unsupported Format", "Please save as .txt or .docx")
                return
            messagebox.showinfo("Export Successful", f"File saved: {filepath}")
        except Exception as e:
            messagebox.showerror("Save Failed", f"Could not save file:\n{e}")

    def toggle_dark_mode(self):
        if not self.dark_mode:
            self.root.configure(bg="#222")
            self.header.configure(bg="#222", fg="white")
            self.tooltip.configure(bg="#555", fg="white")
            self.text_display.configure(bg="#333", fg="#eee", insertbackground="white",
                                        selectbackground="#555")
            self.suggestions_list.configure(bg="#333", fg="white")
            self.upload_btn.configure(bg="#555", fg="white")
            self.auto_correct_btn.configure(bg="#555", fg="white")
            self.export_btn.configure(bg="#555", fg="white")
            self.dark_mode_btn.configure(bg="#777", fg="white")
            self.dark_mode = True
        else:
            self.root.configure(bg="#f2f2f2")
            self.header.configure(bg="#f2f2f2", fg="#333")
            self.tooltip.configure(bg="yellow", fg="black")
            self.text_display.configure(bg="white", fg="black", insertbackground="black",
                                        selectbackground="#cce6ff")
            self.suggestions_list.configure(bg="white", fg="black")
            self.upload_btn.configure(bg="#4CAF50", fg="white")
            self.auto_correct_btn.configure(bg="#2196F3", fg="white")
            self.export_btn.configure(bg="#FF9800", fg="white")
            self.dark_mode_btn.configure(bg="#555", fg="white")
            self.dark_mode = False


if __name__ == "__main__":
    root = tk.Tk()
    app = GrammarSpellChecker(root)
    root.mainloop()
